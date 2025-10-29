import os
import re
import uuid
from typing import List, Optional, Tuple
from pathlib import Path

import chromadb
from chromadb.config import Settings
import openai
from sentence_transformers import SentenceTransformer
from sqlalchemy.orm import Session
import PyPDF2
from docx import Document
try:
    import tiktoken  # optional; may require network on first use
except Exception:  # pragma: no cover
    tiktoken = None

from models import KnowledgeDocument, DocumentChunk, FAQ
from config import settings

class RAGService:
    def __init__(self):
      
        self.chroma_client = chromadb.PersistentClient(
            path="./chroma_db",
            settings=Settings(allow_reset=True)
        )
        
   
        self.collection = self.chroma_client.get_or_create_collection(
            name="knowledge_base",
            metadata={"hnsw:space": "cosine"}
        )
        
        
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        # Optional tokenizer: fall back to word-count if unavailable or offline
        self.tokenizer = None
        if tiktoken is not None:
            try:
                # Prefer a generic encoding; may still try to fetch on first run
                self.tokenizer = tiktoken.get_encoding("cl100k_base")
            except Exception:
                self.tokenizer = None
        
        
        self.kb_triggers = [
            "company", "business", "service", "product", "about us", "what do you do",
            "services", "products", "pricing", "cost", "price", "contact", "location",
            "address", "phone", "email", "team", "staff", "experience", "history",
            "founded", "established", "mission", "vision", "values", "policy",
            "procedure", "process", "how to", "documentation", "manual", "guide",
            "faq", "frequently asked", "help", "support", "technical", "specification"
        ]
    
    def should_use_knowledge_base(self, query: str) -> bool:
        """
        Determine if a query should use the knowledge base based on content analysis.
        Returns False for greetings and casual conversation.
        """
        query_lower = query.lower().strip()
        
        # Always skip KB for casual/personal interactions
        casual_patterns = [
            r'^(hi|hello|hey|good morning|good afternoon|good evening|greetings)!?$',
            r'^(how are you|how\'s it going|what\'s up)!?$',
            r'^(thanks|thank you|bye|goodbye|see you)!?$',
            r'^(yes|no|ok|okay|sure|fine)!?$',
            r'.*\b(critical analysis|analyze me|tell me about myself|personal|bro|dude)\b.*'
        ]
        
        for pattern in casual_patterns:
            if re.match(pattern, query_lower):
                return False
        
        # Look for specific company/business related triggers
        company_triggers = [
            "dipietro", "audiology", "hearing", "appointment", "schedule", "clinic",
            "office hours", "location", "address", "phone number", "contact info",
            "services offered", "pricing", "cost", "insurance", "payment",
            "hearing test", "hearing aid", "doctor", "audiologist"
        ]
        
        for trigger in company_triggers:
            if trigger in query_lower:
                return True
                
        # Only use KB for longer, specific business questions
        if (len(query.split()) > 8 and 
            any(word in query_lower for word in ["appointment", "service", "clinic", "office", "schedule", "available"])):
            return True
            
        return False
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text from various file formats."""
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_extension == '.docx':
            return self._extract_from_docx(file_path)
        elif file_extension == '.txt':
            return self._extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
        return text
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from Word document."""
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from text file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def chunk_text(self, text: str, max_tokens: int = 500) -> List[str]:
        """Split text into chunks based on token count."""
        sentences = re.split(r'[.!?]+', text)
        chunks = []
        current_chunk = ""
        
        def token_len(s: str) -> int:
            if self.tokenizer is not None:
                try:
                    return len(self.tokenizer.encode(s))
                except Exception:
                    pass
            # Fallback: approximate tokens by words
            return max(1, len(s.split()))
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            test_chunk = current_chunk + " " + sentence if current_chunk else sentence
            if token_len(test_chunk) > max_tokens:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                    current_chunk = sentence
                else:
                    
                    words = sentence.split()
                    step = max(1, max_tokens // 4)
                    for i in range(0, len(words), step): 
                        chunk_words = words[i:i + step]
                        chunks.append(" ".join(chunk_words))
            else:
                current_chunk = test_chunk
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def process_document(self, db: Session, file_path: str, filename: str) -> KnowledgeDocument:
        """Process a document and store it in the knowledge base."""
       
        text_content = self.extract_text_from_file(file_path)
        
   
        doc = KnowledgeDocument(
            filename=filename,
            file_path=file_path,
            document_type=Path(file_path).suffix.lower()[1:],  
            processed=False,
            chunk_count=0
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)
        
       
        chunks = self.chunk_text(text_content)
        
       
        chunk_ids = []
        for i, chunk_text in enumerate(chunks):
           
            embedding = self.embedding_model.encode(chunk_text).tolist()
            
        
            vector_id = f"{doc.id}_{i}_{uuid.uuid4().hex[:8]}"
           
            self.collection.add(
                embeddings=[embedding],
                documents=[chunk_text],
                metadatas=[{
                    "document_id": doc.id,
                    "filename": filename,
                    "chunk_index": i
                }],
                ids=[vector_id]
            )
            
           
            chunk = DocumentChunk(
                document_id=doc.id,
                chunk_text=chunk_text,
                chunk_index=i,
                vector_id=vector_id
            )
            db.add(chunk)
            chunk_ids.append(vector_id)
        
       
        doc.processed = True
        doc.chunk_count = len(chunks)
        db.commit()
        
        return doc
    
    def search_faqs(self, db: Session, query: str) -> List[Tuple[str, float, dict]]:
        """Search FAQs for relevant information using simple text matching."""
        # Always search FAQs regardless of query type - FAQs should be used for all questions
        
        query_lower = query.lower()
        try:
            faqs = db.query(FAQ).all()
        except Exception as e:
            return []
        
        faq_results = []
        for faq in faqs:
            question_lower = faq.question.lower()
            answer_lower = faq.answer.lower()
            
            # Calculate relevance score with better matching
            score = 0
            query_words = query_lower.split()
            
            # Exact question match gets highest score
            if query_lower in question_lower:
                score = 100  # Very high score for exact match
            else:
                # Check for word matches with better scoring
                for word in query_words:
                    if len(word) > 2:  # Only consider words longer than 2 characters
                        if word in question_lower:
                            score += 3  # Higher weight for question matches
                        if word in answer_lower:
                            score += 1  # Lower weight for answer matches
                
                # Check for partial matches
                for word in query_words:
                    if len(word) > 3:  # Only for longer words
                        if any(word in q_word for q_word in question_lower.split()):
                            score += 2
                        if any(word in a_word for a_word in answer_lower.split()):
                            score += 1
            
            # Normalize score (0-1 range)
            if score > 0:
                if score >= 100:  # Exact match
                    distance = 0.0
                else:
                    normalized_score = min(1.0, score / max(1, len(query_words) * 3))
                    distance = 1.0 - normalized_score
                
                faq_text = f"Q: {faq.question}\nA: {faq.answer}"
                faq_results.append((
                    faq_text, 
                    distance, 
                    {"source": "faq", "faq_id": faq.id, "question": faq.question}
                ))
        
        # Sort by distance (lower is better) and return top results
        faq_results.sort(key=lambda x: x[1])
        return faq_results[:3]  # Return top 3 FAQ matches

    def search_knowledge_base(self, query: str, top_k: int = 3) -> List[Tuple[str, float, dict]]:
        """Search the knowledge base for relevant information."""
        if not self.should_use_knowledge_base(query):
            return []
        
        
        query_embedding = self.embedding_model.encode(query).tolist()
        
    
        results = self.collection.query(
            query_embeddings=[query_embedding],
            n_results=top_k,
            include=["documents", "metadatas", "distances"]
        )
        
       
        search_results = []
        if results['documents'][0]:
            for i in range(len(results['documents'][0])):
                doc_text = results['documents'][0][i]
                metadata = results['metadatas'][0][i]
                distance = results['distances'][0][i]
                search_results.append((doc_text, distance, metadata))
        
        return search_results
    
    def generate_rag_response(self, query: str, system_prompt: str, db: Session, history: list[dict] | None = None, messaging_config: dict | None = None) -> Tuple[str, bool]:
        """
        Generate a response based on messaging configuration. Use KB when relevant; otherwise regular chat.
        Returns (response, used_kb)
        System prompt takes ABSOLUTE PRIORITY over all other instructions.
        """
        # Default messaging config if not provided
        if messaging_config is None:
            messaging_config = {
                'ai_model': 'gpt-4o',
                'conversational': True,
                'strict_faq': True,
                'response_length': 'Medium',
                'welcome_message': 'Hey there, how can I help you?',
                'server_error_message': 'Apologies, there seems to be a server error.'
            }
        def contextualize_query(q: str, hist: list[dict] | None) -> str:
            q_stripped = (q or "").strip().lower()
            short_ack = {"yes", "yeah", "yup", "y", "no", "nope", "n", "ok", "okay", "sure", "fine", "thanks", "thank you"}
            if len(q_stripped.split()) <= 3 or q_stripped in short_ack:
                if hist:
                    for m in reversed(hist):
                        if m.get("role") == "assistant":
                            prev = m.get("content", "").strip()
                            if prev:
                                return (
                                    f"The user replied: '{q}'. This follows your previous message: '{prev}'. "
                                    f"Respond briefly and take the next helpful step."
                                )
            return q

        q = contextualize_query(query, history)

        # Always search FAQs first - they should be the primary source for all questions
        faq_results = self.search_faqs(db, q)
        
        # If we have good FAQ results, use them
        if faq_results and faq_results[0][1] < 0.5:  # Good FAQ match (distance < 0.5)
            all_results = faq_results
        else:
            # If no good FAQ matches, check if we should use knowledge base
            if self.should_use_knowledge_base(q):
                kb_results = self.search_knowledge_base(q)
                all_results = faq_results + kb_results
                all_results.sort(key=lambda x: x[1])  # Sort by distance (lower is better)
            else:
                # No good FAQ matches and KB not needed, use regular chat
                all_results = []
        
        if not all_results:
            # SYSTEM PROMPT IS ABSOLUTE - use it exactly as provided without modification
            # Only add minimal context about no KB results if needed
            messages: list[dict] = [{"role": "system", "content": system_prompt}]
            if history:
                messages.extend(history)
            messages.append({"role": "user", "content": q})

            client = openai.OpenAI(api_key=settings.OPENAI_API_KEY)
            response = client.chat.completions.create(
                model=messaging_config.get('ai_model', settings.OPENAI_MODEL),
                temperature=0.3,  # Lower temperature for more consistent adherence to instructions
                max_tokens=800,  # Generous limit - let system prompt control length
                messages=messages,
            )
            return response.choices[0].message.content, False

        # Build KB context
        context_parts: list[str] = []
        for doc_text, score, metadata in all_results:
            if score < 0.7:
                if metadata.get('source') == 'faq':
                    context_parts.append(f"FAQ: {doc_text}")
                else:
                    context_parts.append(f"From {metadata.get('filename', 'document')}: {doc_text}")

        if not context_parts:
            # SYSTEM PROMPT IS ABSOLUTE - use it exactly as provided
            messages: list[dict] = [{"role": "system", "content": system_prompt}]
            if history:
                messages.extend(history)
            messages.append({"role": "user", "content": q})

            client = openai.OpenAI(api_key=settings.OPENAI_API_KEY)
            response = client.chat.completions.create(
                model=messaging_config.get('ai_model', settings.OPENAI_MODEL),
                temperature=0.3,  # Lower temperature for more consistent adherence to instructions
                max_tokens=800,  # Generous limit - let system prompt control length
                messages=messages,
            )
            return response.choices[0].message.content, False

        context = "\n\n".join(context_parts)
        
        # CRITICAL: System prompt comes FIRST and is ABSOLUTE
        # KB context is supplementary information, NOT override instructions
        rag_system_prompt = (
            f"{system_prompt}\n\n"
            f"===== RELEVANT KNOWLEDGE BASE INFORMATION =====\n"
            f"{context}\n"
            f"===== END OF KNOWLEDGE BASE INFORMATION =====\n\n"
            f"CRITICAL REMINDER: Follow ALL instructions from your core identity above. "
            f"Use the knowledge base information to enhance your answers, but maintain your communication style, "
            f"response length, and tone exactly as specified in your core identity. "
            f"The knowledge base is supplementary data only - your personality and response rules are ABSOLUTE."
        )

        messages = [{"role": "system", "content": rag_system_prompt}]
        if history:
            messages.extend(history)
        messages.append({"role": "user", "content": q})

        client = openai.OpenAI(api_key=settings.OPENAI_API_KEY)
        response = client.chat.completions.create(
            model=messaging_config.get('ai_model', settings.OPENAI_MODEL),
            temperature=0.3,  # Lower temperature for strict adherence to system prompt
            max_tokens=800,  # Generous limit - let system prompt control the length
            messages=messages,
        )
        return response.choices[0].message.content, True
    
    def delete_document(self, db: Session, document_id: int) -> bool:
        """Delete a document and its chunks from both DB and vector store."""
        
        chunks = db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id).all()
        
    
        vector_ids = [chunk.vector_id for chunk in chunks]
        if vector_ids:
            self.collection.delete(ids=vector_ids)
        
     
        for chunk in chunks:
            db.delete(chunk)
        
       
        doc = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == document_id).first()
        if doc:
            db.delete(doc)
            db.commit()
            return True
        
        return False
    
    def list_documents(self, db: Session) -> List[KnowledgeDocument]:
        """List all documents in the knowledge base."""
        return db.query(KnowledgeDocument).order_by(KnowledgeDocument.upload_date.desc()).all()
    

